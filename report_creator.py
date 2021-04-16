from docxtpl import DocxTemplate, R, Listing

# pylint: disable=E0611
# pylint: disable=E1101
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_LINE_SPACING

from PIL import Image

import os

from calc_okgt import k_conductors


def decToRoman(num):
    rome_nums = ''
    S1 = {1: 'IVX', 10: 'XLC', 100: 'CDM', 1000: 'M  '}
    for i in (1000, 100, 10, 1):
        if num // i != 0:
            a, b, c = S1[i]
            S = (a, a * 2, a * 3, a + b, b, b + a, b + 2 * a, b + 3 * a, a + c)
            rome_nums += S[num // i - 1]
            num = num - (num // i) * i
    return rome_nums

# Функция выравнивания содержимого таблицы по вертикали
def set_cell_vertical_alignment(cell, align="center"): 
    try:   
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')  
        tcValign.set(qn('w:val'), align)  
        tcPr.append(tcValign)
        return True 
    except:
        #traceback.print_exc()             
        return False

def cell_settings(cell, text, align="center", width=None, font_size=12, font_name="Times"):
    cell.text = text
    align_tp = {"center":WD_ALIGN_PARAGRAPH.CENTER,"right":WD_ALIGN_PARAGRAPH.RIGHT,"left":WD_ALIGN_PARAGRAPH.LEFT }
    cell.paragraphs[0].alignment = align_tp[align]
    set_cell_vertical_alignment(cell, align="center")
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    cell.paragraphs[0].runs[0].font.name = font_name

    if width is not None:
        cell.width = Inches(width)

def rpa_settings_table(doc, tbl_ind, vl_name, ps_name, current_rpa_info, report_setings, font_size=12, font_name="Times"):
    p1=doc.add_paragraph()
    p1.add_run(f'Табл. {tbl_ind} - {vl_name} > {ps_name}')
    p1.runs[0].font.size = Pt(font_size)
    p1.runs[0].font.name = font_name
    pt_f = p1.paragraph_format
    pt_f.first_line_indent = Cm(0.25)
    pt_f.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(10)
    

    rows = len(current_rpa_info["rpa_time_setting"])+3 if report_setings["show_arc_pause"] and\
        current_rpa_info["arc_times"]>0 else len(current_rpa_info["rpa_time_setting"])+2
    cols = 3+current_rpa_info["arc_times"] if report_setings["show_Irpa"] else 2+current_rpa_info["arc_times"] 
    # Создаём шапку таблицы
    
    tbl=doc.add_table(rows=rows, cols=cols, style='Table Grid') # Таблица с размерами и стилем  , style='Table Grid'

    tbl.cell(0, 0).merge(tbl.cell(1, 0)) 

    if report_setings["show_Irpa"]:
        tbl.cell(0, 1).merge(tbl.cell(1, 1))
        tbl.cell(0, 2).merge(tbl.cell(0, cols-1))

    else:
        tbl.cell(0, 1).merge(tbl.cell(0, cols-1))

    if report_setings["show_arc_pause"] and current_rpa_info["arc_times"]>0:
        tbl.cell(rows-1, 0).merge(tbl.cell(rows-1, cols-1-current_rpa_info["arc_times"]))

    cell_settings(tbl.cell(0, 0), 'Ступень защиты ВЛ', align="center", font_size=font_size, font_name=font_name)
    tbl.cell(0, 0).width = Inches(0.2)
    

    shift = 0
    if report_setings["show_Irpa"]:
        cell_settings(tbl.cell(0, 1), 'Уставка по току Iкз, кА', align="center", font_size=font_size, font_name=font_name)
        tbl.cell(0, 1).width = Inches(0.2) 
        shift = 1

    cell_settings(tbl.cell(0, 1+shift), 'Время отключения однофазного замыкания на землю Тср.з., с', align="center", font_size=font_size, font_name=font_name)
    tbl.cell(0, 1+shift).width = Inches(0.2) 

    cell_settings(tbl.cell(1, 1+shift), 'Первичное', align="center", font_size=font_size, font_name=font_name)
     
    
    for i in range(current_rpa_info["arc_times"]):
        cell_settings(tbl.cell(1, 2+shift+i), f'АПВ№{i+1}', align="center", font_size=font_size, font_name=font_name)

    if report_setings["show_arc_pause"] and current_rpa_info["arc_times"]>0:
        cell_settings(tbl.cell(rows-1, 0), 'Длительность паузы между отключением КЗ и срабатыванием АПВ, с', align="center", font_size=font_size, font_name=font_name)
        

        for i, val in enumerate(current_rpa_info["arc_pause"]):
            ind = cols-current_rpa_info["arc_times"]+i
            cell_settings(tbl.cell(rows-1, ind), str(round(val,3)), align="center", font_size=font_size, font_name=font_name)

    
    time_setting = current_rpa_info["rpa_time_setting"]

    sort_ind = sorted([i for i in range(len(time_setting))], key=lambda x:time_setting[x])

    rpa_time_setting = [time_setting[i] for i in sort_ind]
    rpa_I_setting = [current_rpa_info["rpa_I_setting"][i] for i in sort_ind]
    arc_setting = [[current_rpa_info["arc_setting"][j][i] for i in sort_ind] for j in range(len(current_rpa_info["arc_setting"]))]

    if report_setings["show_Irpa"]:
        for i, val in enumerate(rpa_I_setting):
            ind = i+2
            cell_settings(tbl.cell(ind, 1), str(round(val,3)), align="center", font_size=font_size, font_name=font_name)
            

    for i, val in enumerate(rpa_time_setting):
        ind = i+2
        cell_settings(tbl.cell(ind, 1+shift), str(round(val,3)), align="center", font_size=font_size, font_name=font_name)
        cell_settings(tbl.cell(ind, 0), decToRoman(i+1), align="center", font_size=font_size, font_name=font_name)

    for j, lst in enumerate(arc_setting):
        c = j+2+shift
        for i, (val1,val2) in enumerate(zip(rpa_time_setting,lst)):
            ind = i+2
            current_time_arc = round(val1-val2,3) if val1-val2>0 else 0
            cell_settings(tbl.cell(ind, c), str(current_time_arc), align="center", font_size=font_size, font_name=font_name)

def colis(n1,k1,n2,k2):
    b = n1<=n2<k2<=k1 or n1>=n2>k2>=k1 or n2<=n1<k1<=k2 or n2>=n1>k1>=k2 or\
        n1<n2<k1 or n1>n2>k1 or n1<k2<k1 or n1>k2>k1 or\
        n2<n1<k2 or n2>n1>k2 or n2<k1<k2 or n2>k1>k2
    return b

def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def find_length(st, ed, val, support):
    trig = True
    for i in range(st,ed):
        if val['links'][i][0][2] == support and trig:
            return val["L"][i]
        if val['links'][i][0][3] == support and not trig:
            return val["L"][i]
        trig = not trig
    else:
        return None

def vl_sector_description(sector,val,vl_info):
    st, ed = sector[2:4]
    start = val['links'][st][0]
    end = val['links'][ed-1][0]
    
    Vlname = start[0]
    ps_name = vl_info[Vlname]["PSs"][0]["PS_name"]

    groundwires = []
    for i in vl_info[start[0]]["groundwires"]:
        if i["link_branch"]==start[1] and colis(start[2],end[3],i["supportN"],i["supportK"]):
            groundwires.append(i)
    
    grounded = []
    for i in vl_info[start[0]]["grounded"]:
        if i["resistance"]<30.0 and i["link_branch"]==start[1] and colis(start[2],end[3],i["supportN"],i["supportK"]):
            grounded.append(i)
    
    countercables = []
    for i,(vl_name, branch,_,_) in enumerate(val['links'][st]):
        if i==0:
            countercables += [i for i in vl_info[vl_name]["countercables"] if i["link_branch"]==branch and colis(start[2],end[3],i["supportN"],i["supportK"])]
        else:
            for item in vl_info[vl_name]["countercables"]:
                if item["link_branch"]==branch:
                    N = item["supportN"]
                    K = item["supportK"]
                    N_new, K_new = -1,-1
                    for j in range(st,ed):
                        for link in val['links'][j]:
                            if (vl_name,branch,N) == link[:3]:
                                N_new = val['links'][0][2]
                            if (vl_name,branch,K) == (link[0],link[1],link[3]):
                                K_new = val['links'][0][3]
                    if N_new!=-1 and K_new!=-1:
                        d = {"supportN":N_new, "supportK": K_new, "link_branch":start[1]}
                        if colis(start[2],end[3],d["supportN"],d["supportK"]):
                            countercables.append({kw:(d[kw] if kw in d else v) for kw, v in item.items()})
    
    support_set = set()
    #print((start[2],end[3]))
    support_set.add(start[2])
    support_set.add(end[3])
    for lst in [grounded,groundwires,countercables]:
        for item in lst:
            if start[2]<=item["supportN"]<=end[3] or start[2]>=item["supportN"]>=end[3]:
                support_set.add(item["supportN"])
            if start[2]<=item["supportK"]<=end[3] or start[2]>=item["supportK"]>=end[3]:
                support_set.add(item["supportK"])

    support = sorted(list(support_set),reverse=start[2]>end[3])
    subsectors = [[support[i-1],support[i]] for i in range(1,len(support))]

    #to_ps = set()
    subsectors_info = []
    for N, K in subsectors:
        wires, ground, conter = None, None, None
        for item in groundwires:
            if colis(N,K,item["supportN"],item["supportK"]):
                okgt = item[item["is_okgt"]]
                R = k_conductors[okgt]["R0"]
                W = k_conductors[okgt]["Bsc"] if k_conductors[okgt]["Bsc"] is not None else 'Неизвестно'
                gw = item.get(("groundwire2" if item["is_okgt"]=="groundwire1" else "groundwire1"),None)
                wires = (okgt,R,gw,W)

        for item in grounded:
            if colis(N,K,item["supportN"],item["supportK"]):
                ground = item["resistance"]

        for item in countercables:
            if colis(N,K,item["supportN"],item["supportK"]):
                connect_to_ps = []
                side_to_ps = []
                if item["connect_to_ps"]:
                    for ps_n, ps_l in val['length_to_ps_lst'][start[0]].items():
                        if ps_l.get((start[1],N),float('inf')) == 0:
                            connect_to_ps.append(ps_n)
                            side_to_ps.append(False)
                            #print("Start", N)
                        if ps_l.get((start[1],K),float('inf')) == 0:
                            connect_to_ps.append(ps_n)
                            side_to_ps.append(True)
                            #print("End", K)
                        
                conter = (item["D_countercable"],connect_to_ps,side_to_ps)  

            
        N_l, K_l = find_length(st, ed, val, N), find_length(st, ed, val, K)
        #print(N,K,wires, ground, conter,N_l, K_l)
        subsectors_info.append((N,K,wires, ground, conter,N_l, K_l))

    return Vlname, ps_name, start, subsectors_info

def conduct_sector_discription(sector,val,branch,okgt_info):
    st, ed = sector[2:4]
    for okgt_sector in okgt_info[branch]:
        if okgt_sector["name"] == sector[0]:
            break

    R = k_conductors[okgt_sector["groundwire"]]["R0"]
    W = k_conductors[okgt_sector["groundwire"]]["Bsc"]

    return val["L"][st],\
        val["L"][ed-1],\
        R,\
        W,\
        okgt_sector["point_grounded"],\
        okgt_sector.get("point_resistance"),\
        okgt_sector["countercable"],\
        okgt_sector.get("D_countercable"),\
        okgt_sector["groundwire"]


def description_settings(doc, okgt_info, vl_info, calc_results, font_size=12, font_name="Times"):
    tbl=doc.add_table(rows=0, cols=5, style='Table Grid')
    for (n,k), val in calc_results.items():
        row_cells = tbl.add_row().cells
        row_cells[0].merge(row_cells[1])
        row_cells[2].merge(row_cells[4])

        cell_settings(row_cells[0], 'Наименование ветви ОКГТ', align="left", font_size=font_size, font_name=font_name)
        cell_settings(row_cells[2], f"{n} - {k}", align="left", font_size=font_size, font_name=font_name)

        for sector in val["sectors"]:
            if sector[1] == "VL":
                
                Vlname, ps_name, start, subsectors_info = vl_sector_description(sector,val,vl_info)

                if subsectors_info:
                    row_cells = tbl.add_row().cells
                    row_cells[0].merge(row_cells[1])
                    row_cells[2].merge(row_cells[4])

                    cell_settings(row_cells[0], 'Наименование участка ОКГТ', align="left", font_size=font_size, font_name=font_name)
                    cell_settings(row_cells[2], sector[0], align="left", font_size=font_size, font_name=font_name)

                    row_cells = tbl.add_row().cells
                    row_cells[0].merge(row_cells[1])

                    cell_settings(row_cells[0], 'Наименование субучастков ОКГТ', align="center", font_size=font_size, font_name=font_name)
                    cell_settings(row_cells[2], 'Сопротивление постоянному току не более, Ом/км', align="center", font_size=font_size, font_name=font_name) 
                    cell_settings(row_cells[3], 'Допустимый тепловой импульс не менее, кА\u00B2·c', align="center", font_size=font_size, font_name=font_name) 
                    cell_settings(row_cells[4], 'Дополнительные мероприятия', align="center", font_size=font_size, font_name=font_name)

                    for item in subsectors_info:
                        row_cells1 = tbl.add_row().cells
                        row_cells2 = tbl.add_row().cells
                        row_cells3 = tbl.add_row().cells
                        row_cells4 = tbl.add_row().cells

                        row_cells1[2].merge(row_cells4[2])
                        row_cells1[3].merge(row_cells4[3])
                        row_cells1[4].merge(row_cells4[4])

                        cell_settings(row_cells1[0], 'ВЛ:', align="left", font_size=font_size, font_name=font_name)
                        cell_settings(row_cells2[0], 'Ветвь:', align="left", font_size=font_size, font_name=font_name)
                        cell_settings(row_cells3[0], 'От:', align="left", font_size=font_size, font_name=font_name)
                        cell_settings(row_cells4[0], 'До:', align="left", font_size=font_size, font_name=font_name)

                        cell_settings(row_cells1[1], Vlname, align="left", font_size=font_size, font_name=font_name)
                        cell_settings(row_cells2[1], f'{start[1][0]} - {start[1][1]}', align="left", font_size=font_size, font_name=font_name)

                        fr = round(val['length_to_ps_lst'][Vlname][ps_name][(start[1],item[0])],3)
                        t = round(val['length_to_ps_lst'][Vlname][ps_name][(start[1],item[1])],3)

                        cell_settings(row_cells3[1], f'{fr} км от {ps_name}' , align="left", font_size=font_size, font_name=font_name)
                        cell_settings(row_cells4[1], f'{t} км от {ps_name}', align="left", font_size=font_size, font_name=font_name)

                        #print(item[2])
                        cell_settings(row_cells1[2], str(item[2][1]) , align="center", font_size=font_size, font_name=font_name)
                        cell_settings(row_cells1[3], str(item[2][3]), align="center", font_size=font_size, font_name=font_name)

                        s = []
                        count = 0
                        if item[3] is not None:
                            count += 1
                            s.append(f'{count}. Обеспечить на данном участке сопротивление ЗУ опор не более {round(item[3],2)} Ом')
                            

                        if item[2][2] is not None:
                            count += 1
                            s.append(f'{count}. Подвесить на данном участке второй грозотрос марки {item[2][2]}')
                            
                        
                        if item[4] is not None:
                            count += 1
                            a = f'{count}. ЗУ опор соединить между собой горизонтальным заземлителем из стали круглой диаметром {round(item[4][0],1)} мм'
                            if item[4][1]:
                                a += ' и присоединить к ЗУ %s' % ','.join(item[4][1])
                            s.append(a)

                        advices = ';\n'.join(s)

                        if len(advices)!=0:
                            cell_settings(row_cells1[4], advices+'.', align="left", font_size=font_size, font_name=font_name)
                        else:
                            cell_settings(row_cells1[4], '-', align="center", font_size=font_size, font_name=font_name)
                        
            elif sector[1] == "single_dielectric":
                st, ed = sector[2:4]

                row_cells = tbl.add_row().cells
                row_cells[0].merge(row_cells[1])
                row_cells[2].merge(row_cells[4])

                cell_settings(row_cells[0], 'Наименование участка ОКГТ', align="left", font_size=font_size, font_name=font_name)
                cell_settings(row_cells[2], sector[0], align="left", font_size=font_size, font_name=font_name)

                row_cells = tbl.add_row().cells
                row_cells[0].merge(row_cells[1])

                cell_settings(row_cells[0], 'Наименование субучастков ОКГТ', align="center", font_size=font_size, font_name=font_name)
                cell_settings(row_cells[2], 'Сопротивление постоянному току не более, Ом/км', align="center", font_size=font_size, font_name=font_name) 
                cell_settings(row_cells[3], 'Допустимый тепловой импульс, кА\u00B2·c', align="center", font_size=font_size, font_name=font_name) 
                cell_settings(row_cells[4], 'Дополнительные мероприятия', align="center",font_size=font_size, font_name=font_name)

                row_cells1 = tbl.add_row().cells
                row_cells2 = tbl.add_row().cells

                row_cells1[0].merge(row_cells1[1])
                row_cells1[2].merge(row_cells2[2])
                row_cells1[3].merge(row_cells2[3])
                row_cells1[4].merge(row_cells2[4])

                cell_settings(row_cells1[0], 'Подземный диэлектрический ВОК', align="center", font_size=font_size, font_name=font_name)
                cell_settings(row_cells2[0], 'Длина:', align="left", font_size=font_size, font_name=font_name)
                cell_settings(row_cells2[1], f'{round(abs(val["L"][ed-1]-val["L"][st]),3)} км', align="left", font_size=font_size, font_name=font_name)

                cell_settings(row_cells1[2], '-', align="center", font_size=font_size, font_name=font_name)
                cell_settings(row_cells1[3], '-', align="center", font_size=font_size, font_name=font_name)
                cell_settings(row_cells1[4], '-', align="center", font_size=font_size, font_name=font_name)

            elif sector[1] == "single_conductive":
                

                lng_st,lng_ed,R,W,point_grounded,point_resistance,countercable,D_countercable,_ = conduct_sector_discription(sector,val,(n,k),okgt_info)

                length = round(abs(lng_ed-lng_st),3)
                s = []
                count = 0
                if point_grounded>0:
                    if point_resistance<30.0:
                        count += 1
                        s.append(f'{count}. Обеспечить на данном участке сопротивление ЗУ опор не более {round(point_resistance,2)} Ом')

                if countercable:
                    count += 1
                    s.append(f'{count}. ЗУ опор соединить между собой горизонтальным заземлителем из стали круглой диаметром {round(D_countercable,1)} мм')

                if W is not None:

                    row_cells = tbl.add_row().cells
                    row_cells[0].merge(row_cells[1])
                    row_cells[2].merge(row_cells[4])

                    cell_settings(row_cells[0], 'Наименование участка ОКГТ', align="left", font_size=font_size, font_name=font_name)
                    cell_settings(row_cells[2], sector[0], align="left", font_size=font_size, font_name=font_name)

                    row_cells = tbl.add_row().cells
                    row_cells[0].merge(row_cells[1])

                    cell_settings(row_cells[0], 'Наименование субучастков ОКГТ', align="center", font_size=font_size, font_name=font_name)
                    cell_settings(row_cells[2], 'Сопротивление постоянному току не более, Ом/км', align="center", font_size=font_size, font_name=font_name) 
                    cell_settings(row_cells[3], 'Допустимый тепловой импульс, кА\u00B2·c', align="center", font_size=font_size, font_name=font_name) 
                    cell_settings(row_cells[4], 'Дополнительные мероприятия', align="center",font_size=font_size, font_name=font_name)

                    row_cells1 = tbl.add_row().cells
                    row_cells2 = tbl.add_row().cells

                    row_cells1[0].merge(row_cells1[1])
                    row_cells1[2].merge(row_cells2[2])
                    row_cells1[3].merge(row_cells2[3])
                    row_cells1[4].merge(row_cells2[4])

                    cell_settings(row_cells1[0], 'Электропроводящий ВОК', align="center", font_size=font_size, font_name=font_name)
                    cell_settings(row_cells2[0], 'Длина:', align="left", font_size=font_size, font_name=font_name)
                    cell_settings(row_cells2[1], f'{length} км', align="left", font_size=font_size, font_name=font_name)

                    cell_settings(row_cells1[2], str(R), align="center", font_size=font_size, font_name=font_name)
                    cell_settings(row_cells1[3], str(W) if W is not None else 'Неизвестно', align="center", font_size=font_size, font_name=font_name)


                    advices = ';\n'.join(s)

                    if len(advices)!=0:
                        cell_settings(row_cells1[4], advices+'.', align="left", font_size=font_size, font_name=font_name)
                    else:
                        cell_settings(row_cells1[4], '-', align="center", font_size=font_size, font_name=font_name)
                    
    set_col_widths(tbl, (Cm(2),Cm(2.5),Cm(5),Cm(4),Cm(5.5)))

                
def grounded_wires(doc, calc_results, font_size=12, font_name="Times"):
    tbl=doc.add_table(rows=1, cols=2, style='Table Grid')
    sets = []
    for i in calc_results.values():
        sets+=i['okgt_types']
    recomends = list(set().union(*sets).difference(set([None])))
    recomends = sorted(recomends, key=lambda x: k_conductors[x]["Bsc"])
    cell_settings(tbl.cell(0, 0), 'Допустимый тепловой импульс ОКГТ, кА\u00B2·c', align="center", font_size=font_size, font_name=font_name)
    cell_settings(tbl.cell(0, 1), 'Марка заземляющего проводника', align="center", font_size=font_size, font_name=font_name)
    for item in recomends:
        if k_conductors[item]["Grounded_conductor"] is not None:
            row_cells = tbl.add_row().cells
            cell_settings(row_cells[0], str(int(k_conductors[item]["Bsc"])), align="center", font_size=font_size, font_name=font_name)
            cell_settings(row_cells[1], k_conductors[item]["Grounded_conductor"], align="center", font_size=font_size, font_name=font_name)
            

    


def memorandum(fname, path_midle_files, okgt_info, vl_info, rpa_info, calc_results, report_setings):
    doc = DocxTemplate("docx_templates/memorandum.docx")
    context = { 
        'recipients' : R('\n'.join([i.strip() for i in report_setings['recipients'].split(';')])),
        "project_name" : report_setings['project_name']
    }
    doc.render(context)

    midle_docx_file = os.path.join(path_midle_files,'midle_memorandum.docx')       
    
    doc.save(midle_docx_file)

    document = Document(midle_docx_file)
    #document = Document()

    # Наследуем стиль и изменяем его
    style = document.styles['Normal'] # Берём стиль Нормальный
    f0 = style.font # Переменная для изменения параметров стиля
    f0.name = 'Times New Roman' # Шрифт
    f0.size = Pt(14) # Размер шрифта

    description_settings(document,okgt_info, vl_info, calc_results, font_name='Times New Roman')

    p1=document.add_paragraph()
    p1.add_run(f'2) Параметры работы РЗ приняты согласно Служебной записке ОРЗА таблицах 2-{len(rpa_info)+1}')
    p1.runs[0].font.size = Pt(14) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(20)

    for i, ((vl_name,ps_name), rpa) in enumerate(rpa_info.items()):
        rpa_settings_table(document, i+2, vl_name, ps_name, rpa, report_setings, font_name='Times New Roman')

    p1=document.add_paragraph()
    p1.add_run(f'3) Для присоединения ОКГТ к опорам ВЛ использовать провода, указанные в таблице {len(rpa_info)+2}.')
    p1.runs[0].font.size = Pt(14) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1)
    pt_f.space_after = Pt(10)
    pt_f.space_before = Pt(20)

    p1=document.add_paragraph()
    p1.add_run(f'Табл. {len(rpa_info)+2} – Рекомендации по выбору заземляющего проводника')
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(0.25)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(0)

    grounded_wires(document, calc_results)

    p1=document.add_paragraph()
    p1.add_run('4) Стойкость ОКГТ к воздействию импульса грозового разряда:')
    p1.runs[0].font.size = Pt(14) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1)
    pt_f.space_after = Pt(5)
    pt_f.space_before = Pt(20)

    p1=document.add_paragraph()
    p1.add_run('- амплитуда тока молнии не менее 100 кА;')
    p1.runs[0].font.size = Pt(14) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(0)

    p1=document.add_paragraph()
    p1.add_run('- заряд молнии не менее 50 Кл.')
    p1.runs[0].font.size = Pt(14) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent= Cm(1.25)
    pt_f.space_after = Pt(15)
    pt_f.space_before = Pt(0)


    p1=document.add_paragraph()
    p1.add_run(('Начальник ОУКЭ\t' if not report_setings["department_boss_type"] else 'Зам. начальника ОУКЭ')+\
        '\t\t\t\t\t\t\t'+report_setings["department_boss_name"])
    p1.runs[0].font.size = Pt(14) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(0)
    pt_f.space_after = Pt(25)
    pt_f.space_before = Pt(30)


    p1=document.add_paragraph()
    p1.add_run(f'Зав. гр. ТВН ОУКЭ\t\t\t\t\t\t\t{report_setings["group_boss_name"]}')
    p1.runs[0].font.size = Pt(14) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(0)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(25)



    document.save(fname)
    print('memorandum was made')

def length_subsectors(length,tp):
    lst = []
    l_s, t_s = length[0], tp[0]
    for i in range(1,len(length)):
        if t_s!=tp[i]:
            lst.append((t_s,abs(length[i-1]-l_s)))
            t_s = tp[i]
            l_s = length[i]

    else:
        lst.append((t_s,abs(length[i]-l_s)))

    return lst

def explanatory(fname, path_midle_files, okgt_info, vl_info, rpa_info, report_setings, calc_results, sectorsFig, rpa_liks):
    document = Document("docx_templates/explanatory_note.docx")

    style = document.styles['Normal'] # Берём стиль Нормальный
    f0 = style.font # Переменная для изменения параметров стиля
    f0.name = 'Arial' # Шрифт
    f0.size = Pt(12) # Размер шрифта
    document.styles['Normal'].paragraph_format.line_spacing = 1

    tables_count = len(rpa_info)

    p1=document.add_paragraph()
    p1.add_run(f'Исходные данные для расчета допустимого теплового импульса представлены в табл.1-{tables_count} и на рис.1-{tables_count}.')
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(5)
    pt_f.space_before = Pt(0)

    for i, ((vl_name,ps_name), rpa) in enumerate(rpa_info.items()):
        rpa_settings_table(document, i+1, vl_name, ps_name, rpa, report_setings, font_size=10, font_name='Arial')


    p1=document.add_paragraph()
    p1.add_run(('Распределение токов ОКЗ по ВЛ представлено в виде кривых, отображающих составляющие '
        'тока ОКЗ со стороны питающих концов ВЛ и кривой суммарного тока ОКЗ. Кривые распределения тока '
        'строятся на основании расчета тока ОКЗ при замыкании в заданных точках. '
        f'Кривые распределения тока ОКЗ по ВЛ представлены на рис.1-{tables_count}.'))
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(5)
    pt_f.space_before = Pt(12)

    for i, val in enumerate(rpa_liks.values()):
        vl_name = val['vl_combo'].currentText()
        ps_name = val['ps_combo'].currentText()
        midle_img = os.path.join(path_midle_files,f'figures/ssc_{i}.jpg')
        val['figure'].set_size_inches(17, 5,forward=True) # Изменяем размер сохраняемого графика
        val['figure'].savefig(midle_img, format='jpg', dpi=100) # Cохраняем графики
        im = Image.open(midle_img)
        width, height = im.size
        im.crop((130,0,width-130,height)).save(midle_img)
        document.add_picture(midle_img, width=Inches(6.5))

        p1=document.add_paragraph()
        p1.add_run(f'Рис.{i+1} – Распределение тока ОКЗ по проводам {vl_name} от {ps_name}')
        p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
        pt_f = p1.paragraph_format
        pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt_f.first_line_indent = Cm(0)
        #pt_f.left_indent = Cm(0)
        pt_f.space_after = Pt(0)
        pt_f.space_before = Pt(0)

    p1=document.add_paragraph(style=document.styles['Heading 2'])
    p1.add_run('Определение теплового импульса в ОКГТ')

    current_fig = tables_count+1
    fig_count = len(sectorsFig)+tables_count

    p1=document.add_paragraph()
    p1.add_run(('По результатам расчета строятся кривые распределения теплового импульса '
        f'в ОКГТ (W=f(№опор)), приведенные на рис. {current_fig}-{fig_count}. В скобках указываются параметры '
        'про-веряемого ОКГТ: сопротивление постоянному току RDC и тепловой импульс W.'))
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(5)
    pt_f.space_before = Pt(0)

    i=-1
    for (n,k),val in calc_results.items():
        for sector in val["sectors"]:
            if sector[1] != "single_dielectric":
                i+=1
                midle_img = os.path.join(path_midle_files,f'figures/w_{i}.jpg')
                sectorsFig[sector][0].set_size_inches(17, 5,forward=True) 
                sectorsFig[sector][0].savefig(midle_img, format='jpg', dpi=100) 
                im = Image.open(midle_img)
                width, height = im.size
                im.crop((130,0,width-130,height)).save(midle_img)
                document.add_picture(midle_img, width=Inches(6.5))

                st, ed = sector[2:4]
                fig_title = f'Рис.{current_fig+i} – Распределение теплового импульса в ОКГТ на ветви {n} - {k}, участке {sector[0]} '
                s = []
                for okgt_tp, length in length_subsectors(val["L"][st:ed],val["conductor"][st:ed]):
                    if k_conductors.get(okgt_tp,{}).get("Bsc",None) is not None:
                        s.append(f'{round(length,3)} км - Rdc={round(k_conductors[okgt_tp]["R0"],3)} Ом/км, W={int(k_conductors[okgt_tp]["Bsc"])} кА\u00B2·c')
                
                if s:
                   fig_title += '(' + '; '.join(s) + ')'

                p1=document.add_paragraph()
                p1.add_run(fig_title)
                p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
                pt_f = p1.paragraph_format
                pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pt_f.first_line_indent = Cm(0)
                #pt_f.left_indent = Cm(0)
                pt_f.space_after = Pt(0)
                pt_f.space_before = Pt(0)

    
    p1=document.add_paragraph(style=document.styles['Heading 2'])
    p1.add_run('Рекомендации по выбору ОКГТ по условию допустимого теплового импульса')

    p1=document.add_paragraph()
    p1.add_run(('В результате выполненных расчётов определены области применения ОКГТ '
        'с допустимым тепловым импульсом не менее и сопротивлением постоянному току не более указанных в табл.2.'))
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(6)
    pt_f.space_before = Pt(0)

    p1=document.add_paragraph()
    p1.add_run(f'Табл. {tables_count+1} – Рекомендации по выбору ОКГТ')
    p1.runs[0].font.size = Pt(10) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(0.25)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(0)

    description_settings(document,okgt_info, vl_info, calc_results, font_size=10, font_name='Arial')


    p1=document.add_paragraph(style=document.styles['Heading 2'])
    p1.add_run('Рекомендации по выбору ОКГТ по условию устойчивости к прямому удару молнии')

    p1=document.add_paragraph()
    p1.add_run(('Выбор типа ОКГТ по условию устойчивости к прямому удару молнии произво-дится в '
        'соответствии с «Инструкцией по устройству молниезащиты зданий, '
        'сооружений и промышленных коммуникаций», СО 153-34.21.122-2003, РФ.'))
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(0)

    p1=document.add_paragraph()
    p1.add_run('Принимаем III уровень защиты от прямого удара молнии с коэффициентом надежности равным 0,90, имеющим следующие требования:')
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(0)

    p1=document.add_paragraph()
    p1.add_run('- амплитуда тока молнии не менее 100 кА;')
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(0)

    p1=document.add_paragraph()
    p1.add_run('- заряд импульса молнии не менее 50 Кл.')
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(0)

    p1=document.add_paragraph(style=document.styles['Heading 2'])
    p1.add_run('Рекомендации по выбору заземляющего проводника для присоединения ОКГТ к опорам ВЛ')

    p1=document.add_paragraph()
    p1.add_run(('Заземляющий проводник для присоединения ОКГТ к опорам ВЛ должен вы-держивать тепловой импульс, '
        'создаваемый частью тока ОКЗ стекающей с опоры в ОКГТ. При расчете теплового импульса в ОКГТ учитывается, '
        'что ток растекается по ОКГТ в обе стороны от опоры. Следовательно, заземляющий проводник должен вы-держивать тепловой импульс, '
        'превышающий допустимый тепловой импульс ОКГТ, для присоединения которого он служит.'))
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(0)

    p1=document.add_paragraph()
    p1.add_run(('Определение допустимого теплового импульса для заземляющего проводника производится '
        'исходя из допустимой величины тока КЗ в соответствии с «Методиче-скими указаниями по расчету '
        'термической устойчивости грозозащитных тросов воз-душных линий электропередачи», №5290тм-т1, «Энергосетьпроект», Киев, 1976 г.'))
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(0)

    p1=document.add_paragraph()
    p1.add_run((f'Рекомендуемые марки заземляющего проводника для заземления ОКГТ при-ведены в табл. {tables_count+2}.'))
    p1.runs[0].font.size = Pt(12) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(1.25)
    #pt_f.left_indent = Cm(0)
    pt_f.space_after = Pt(6)
    pt_f.space_before = Pt(0)

    p1=document.add_paragraph()
    p1.add_run(f'Табл. {tables_count+2} – Рекомендации по выбору заземляющего проводника')
    p1.runs[0].font.size = Pt(10) #Меняем размер шрифта параграфа
    pt_f = p1.paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.first_line_indent = Cm(0.25)
    pt_f.space_after = Pt(0)
    pt_f.space_before = Pt(0)

    grounded_wires(document, calc_results)
    

    document.save(fname)
    print('explanatory was made')

    